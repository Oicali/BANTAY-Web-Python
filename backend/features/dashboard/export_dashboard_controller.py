# backend\features\dashboard\export_dashboard_controller.py

"""
Generates a PDF report from crime dashboard data.
Builds a .docx in-memory using python-docx, then converts to PDF via LibreOffice.
"""

import io
import base64
import subprocess
import tempfile
import os
from datetime import datetime, timezone, timedelta

from fastapi import Request
from fastapi.responses import Response, JSONResponse
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from shared.utils.audit_logger import log_audit, get_client_ip

# ── Constants ─────────────────────────────────────────────────────────────────
CRIME_DISPLAY = {
    "MURDER":               "Murder",
    "HOMICIDE":             "Homicide",
    "PHYSICAL INJURY":      "Physical Injury",
    "RAPE":                 "Rape",
    "ROBBERY":              "Robbery",
    "THEFT":                "Theft",
    "CARNAPPING - MC":      "Carnapping - MC",
    "CARNAPPING - MV":      "Carnapping - MV",
    "SPECIAL COMPLEX CRIME":"Special Complex Crime",
}

# Colors (hex strings, no #)
DARK  = "1E3A5F"
LIGHT = "D9E4F0"
WHITE = "FFFFFF"
GRAY  = "F3F4F6"

# EMU dimensions for full-width chart images
FULL_W_EMU = 6680400
FULL_H_EMU = 1854000


# ── Pure helpers ──────────────────────────────────────────────────────────────
def pct(n: int, d: int) -> str:
    return f"{(n / d * 100):.1f}" if d else "0.0"


def fmt_date_iso(iso: str) -> str:
    if not iso:
        return ""
    y, m, d = iso.split("-")
    return f"{d}/{m}/{y}"


def hex_to_rgb(h: str) -> RGBColor:
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


# ── Low-level XML helpers ─────────────────────────────────────────────────────
def _set_cell_shading(cell, fill_hex: str):
    """Apply solid background fill to a table cell via raw XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = "CCCCCC"):
    """Apply thin single-line borders to all four sides of a cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _set_cell_width(cell, width_dxa: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_dxa))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_cell_valign(cell, align: str = "center"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _set_table_width(table, width_dxa: int):
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"),    str(width_dxa))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)


def _set_column_widths(table, widths: list[int]):
    tbl = table._tbl
    tblGrid = OxmlElement("w:tblGrid")
    for w in widths:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(w))
        tblGrid.append(gridCol)
    tbl.insert(1, tblGrid)


def _para_border_bottom(para, color: str = DARK, size: int = 6):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(size))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _para_spacing(para, before: int = 0, after: int = 0):
    pPr     = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    if before: spacing.set(qn("w:before"), str(before))
    if after:  spacing.set(qn("w:after"),  str(after))
    pPr.append(spacing)


def _run_font(run, size_half_pt: int, bold: bool = False,
              color_hex: str = "000000", font: str = "Arial"):
    run.bold = bold
    run.font.name  = font
    run.font.size  = Pt(size_half_pt / 2)
    run.font.color.rgb = hex_to_rgb(color_hex)


# ── Cell builders ─────────────────────────────────────────────────────────────
def h_cell(table_row, text: str, width_dxa: int,
           center: bool = False, small: bool = False):
    """Header cell — dark background, white bold text."""
    cell = table_row.add_cell()
    cell.text = ""
    _set_cell_shading(cell, DARK)
    _set_cell_borders(cell)
    _set_cell_margins(cell,
                      *(60, 60, 80, 80) if small else (80, 80, 120, 120))
    _set_cell_width(cell, width_dxa)
    _set_cell_valign(cell, "center")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    _run_font(run, 14 if small else 18, bold=True, color_hex=WHITE)
    return cell


def d_cell(table_row, text, width_dxa: int,
           center: bool = False, bold: bool = False,
           alt: bool = False, shade: bool = False,
           color_hex: str = "000000", small: bool = False):
    """Data cell — white/gray/light-blue background."""
    fill = LIGHT if shade else (GRAY if alt else WHITE)
    cell = table_row.add_cell()
    cell.text = ""
    _set_cell_shading(cell, fill)
    _set_cell_borders(cell)
    _set_cell_margins(cell,
                      *(60, 60, 80, 80) if small else (80, 80, 120, 120))
    _set_cell_width(cell, width_dxa)
    _set_cell_valign(cell, "center")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text) if text is not None else "")
    _run_font(run, 14 if small else 18, bold=bold, color_hex=color_hex)
    return cell


# ── Paragraph helpers ─────────────────────────────────────────────────────────
def section_heading(doc: Document, text: str):
    p = doc.add_paragraph()
    _para_spacing(p, before=300, after=120)
    _para_border_bottom(p, color=DARK, size=6)
    run = p.add_run(text)
    _run_font(run, 28, bold=True, color_hex=DARK)
    return p


def body_text(doc: Document, text: str):
    p = doc.add_paragraph()
    _para_spacing(p, after=80)
    run = p.add_run(text)
    _run_font(run, 20)
    return p


def spacer(doc: Document, before: int = 120):
    p = doc.add_paragraph()
    _para_spacing(p, before=before)
    return p


def image_block(doc: Document, b64: str):
    if not b64:
        return
    data = base64.b64decode(b64)
    stream = io.BytesIO(data)
    p = doc.add_paragraph()
    _para_spacing(p, before=80, after=120)
    run = p.add_run()
    run.add_picture(stream, width=Emu(FULL_W_EMU), height=Emu(FULL_H_EMU))


# ── Table builders ────────────────────────────────────────────────────────────
def build_index_crime_table(doc: Document, summary: list):
    COL = [2766, 1100, 1100, 1100, 1300, 1100, 1000]
    TWIDTH = sum(COL)

    totals = {
        "total":   sum(r["total"] for r in summary),
        "cleared": sum(r["cleared"] for r in summary),
        "solved":  sum(r["solved"] for r in summary),
        "ui":      sum(r["underInvestigation"] for r in summary),
    }

    table = doc.add_table(rows=0, cols=len(COL))
    _set_table_width(table, TWIDTH)
    _set_column_widths(table, COL)

    # Header
    hr = table.add_row()
    h_cell(hr, "Index Crime",  COL[0])
    h_cell(hr, "Total",        COL[1], center=True)
    h_cell(hr, "Cleared",      COL[2], center=True)
    h_cell(hr, "Solved",       COL[3], center=True)
    h_cell(hr, "Under Inv.",   COL[4], center=True)
    h_cell(hr, "CCE %",        COL[5], center=True)
    h_cell(hr, "CSE %",        COL[6], center=True)

    for i, row in enumerate(summary):
        cce = pct(row["cleared"] + row["solved"], row["total"])
        cse = pct(row["solved"], row["total"])
        alt = i % 2 == 1
        dr = table.add_row()
        d_cell(dr, CRIME_DISPLAY.get(row["crime"], row["crime"]), COL[0], alt=alt)
        d_cell(dr, row["total"],              COL[1], center=True, alt=alt)
        d_cell(dr, row["cleared"],            COL[2], center=True, alt=alt, color_hex="1d4ed8")
        d_cell(dr, row["solved"],             COL[3], center=True, alt=alt, color_hex="15803d")
        d_cell(dr, row["underInvestigation"], COL[4], center=True, alt=alt, color_hex="b45309")
        d_cell(dr, f"{cce}%",                COL[5], center=True, alt=alt)
        d_cell(dr, f"{cse}%",                COL[6], center=True, alt=alt)

    # Totals row
    tr = table.add_row()
    d_cell(tr, "TOTAL",                                              COL[0], bold=True, shade=True)
    d_cell(tr, totals["total"],                                      COL[1], center=True, bold=True, shade=True)
    d_cell(tr, totals["cleared"],                                    COL[2], center=True, bold=True, shade=True)
    d_cell(tr, totals["solved"],                                     COL[3], center=True, bold=True, shade=True)
    d_cell(tr, totals["ui"],                                         COL[4], center=True, bold=True, shade=True)
    d_cell(tr, f"{pct(totals['cleared']+totals['solved'],totals['total'])}%",
                                                                     COL[5], center=True, bold=True, shade=True)
    d_cell(tr, f"{pct(totals['solved'], totals['total'])}%",         COL[6], center=True, bold=True, shade=True)

    return table


def build_summary_stat_table(doc: Document, totals: dict, cce: str, cse: str):
    stats = [
        ("Total Incidents",    str(totals["total"])),
        ("Cleared",            str(totals["cleared"])),
        ("Solved",             str(totals["solved"])),
        ("Under Investigation",str(totals["ui"])),
        ("CCE %",              f"{cce}%"),
        ("CSE %",              f"{cse}%"),
    ]
    COL = [3200, 3200]
    table = doc.add_table(rows=0, cols=2)
    _set_table_width(table, sum(COL))
    _set_column_widths(table, COL)
    for i, (label, val) in enumerate(stats):
        alt = i % 2 == 1
        row = table.add_row()
        d_cell(row, label, COL[0], alt=alt)
        d_cell(row, val,   COL[1], bold=True, alt=alt)
    return table


def build_by_day_table(doc: Document, by_day: list):
    COL   = [4200, 2000]
    total = sum(r.get("count", 0) for r in by_day)
    table = doc.add_table(rows=0, cols=2)
    _set_table_width(table, sum(COL))
    _set_column_widths(table, COL)

    hr = table.add_row()
    h_cell(hr, "Day of Week", COL[0])
    h_cell(hr, "Incidents",   COL[1], center=True)

    for i, row in enumerate(by_day):
        dr = table.add_row()
        d_cell(dr, row["day"],   COL[0], alt=i % 2 == 1)
        d_cell(dr, row["count"], COL[1], center=True, bold=True, alt=i % 2 == 1)

    tr = table.add_row()
    d_cell(tr, "TOTAL", COL[0], bold=True, shade=True)
    d_cell(tr, total,   COL[1], center=True, bold=True, shade=True)
    return table


def build_modus_table(doc: Document, modus: list):
    COL   = [2200, 3566, 900]
    table = doc.add_table(rows=0, cols=3)
    _set_table_width(table, sum(COL))
    _set_column_widths(table, COL)

    hr = table.add_row()
    h_cell(hr, "Crime", COL[0])
    h_cell(hr, "Modus", COL[1])
    h_cell(hr, "Count", COL[2], center=True)

    for i, row in enumerate(modus):
        dr = table.add_row()
        d_cell(dr, CRIME_DISPLAY.get(row["crime"], row["crime"]), COL[0], alt=i % 2 == 1)
        d_cell(dr, row["modus"], COL[1], alt=i % 2 == 1)
        d_cell(dr, row["count"], COL[2], center=True, bold=True, alt=i % 2 == 1)
    return table


def build_place_table(doc: Document, place: list):
    COL   = [700, 4966, 1200]
    table = doc.add_table(rows=0, cols=3)
    _set_table_width(table, sum(COL))
    _set_column_widths(table, COL)

    hr = table.add_row()
    h_cell(hr, "#",        COL[0], center=True)
    h_cell(hr, "Location", COL[1])
    h_cell(hr, "Count",    COL[2], center=True)

    for i, row in enumerate(place):
        dr = table.add_row()
        d_cell(dr, i + 1,      COL[0], center=True, alt=i % 2 == 1)
        d_cell(dr, row["place"],COL[1], alt=i % 2 == 1)
        d_cell(dr, row["count"],COL[2], center=True, bold=True, alt=i % 2 == 1)
    return table


def build_barangay_table(doc: Document, barangay: list):
    COL   = [700, 4966, 1200]
    table = doc.add_table(rows=0, cols=3)
    _set_table_width(table, sum(COL))
    _set_column_widths(table, COL)

    hr = table.add_row()
    h_cell(hr, "#",         COL[0], center=True)
    h_cell(hr, "Barangay",  COL[1])
    h_cell(hr, "Incidents", COL[2], center=True)

    for i, row in enumerate(barangay):
        dr = table.add_row()
        d_cell(dr, i + 1,          COL[0], center=True, alt=i % 2 == 1)
        d_cell(dr, row["barangay"],COL[1], alt=i % 2 == 1)
        d_cell(dr, row["count"],   COL[2], center=True, bold=True, alt=i % 2 == 1)
    return table


def build_complete_data_table(doc: Document, complete_data: list):
    """Uses smaller font (7pt = size 14 half-pts) to fit a wide table."""
    COL   = [1800, 1500, 900, 800, 1800, 2666, 1000]
    table = doc.add_table(rows=0, cols=7)
    _set_table_width(table, sum(COL))
    _set_column_widths(table, COL)

    hr = table.add_row()
    h_cell(hr, "Barangay",      COL[0], small=True)
    h_cell(hr, "Type of Place", COL[1], small=True)
    h_cell(hr, "Date",          COL[2], small=True, center=True)
    h_cell(hr, "Time",          COL[3], small=True, center=True)
    h_cell(hr, "Crime Offense", COL[4], small=True)
    h_cell(hr, "Modus",         COL[5], small=True)
    h_cell(hr, "Case Status",   COL[6], small=True)

    for i, row in enumerate(complete_data):
        status_lower = (row.get("caseStatus") or "").lower()
        is_ui      = status_lower not in ("cleared","cce","solved","cse","closed")
        is_solved  = status_lower in ("solved","cse")
        is_cleared = status_lower in ("cleared","cce")
        status_color = ("b45309" if is_ui
                        else "15803d" if is_solved
                        else "1d4ed8" if is_cleared
                        else "000000")
        alt = i % 2 == 1
        dr  = table.add_row()
        d_cell(dr, row.get("barangay",""),                                      COL[0], alt=alt, small=True)
        d_cell(dr, row.get("typeOfPlace",""),                                   COL[1], alt=alt, small=True)
        d_cell(dr, row.get("date",""),                                          COL[2], center=True, alt=alt, small=True)
        d_cell(dr, row.get("time",""),                                          COL[3], center=True, alt=alt, small=True)
        d_cell(dr, CRIME_DISPLAY.get(row.get("crimeOffense",""), row.get("crimeOffense","")),
                                                                                COL[4], alt=alt, small=True)
        d_cell(dr, row.get("modus",""),                                         COL[5], alt=alt, small=True)
        d_cell(dr, row.get("caseStatus",""),                                    COL[6], alt=alt,
               color_hex=status_color, small=True)
    return table


def build_assessment_section(doc: Document, assessment: dict | None,
                              analysis_data: dict | None):
    if not assessment:
        return

    section_heading(doc, "10. AI Crime Assessment")

    # Scope
    if assessment.get("scope"):
        scope = assessment["scope"]
        COL   = [2400, 8066]
        tbl   = doc.add_table(rows=0, cols=2)
        _set_table_width(tbl, sum(COL))
        _set_column_widths(tbl, COL)
        for i, (label, val) in enumerate([
            ("Date Range", scope.get("dateRange") or "-"),
            ("Crime Type", scope.get("crimes")    or "-"),
            ("Barangay",   scope.get("barangays")  or "-"),
        ]):
            r = tbl.add_row()
            d_cell(r, label, COL[0], alt=i % 2 == 1)
            d_cell(r, val,   COL[1], alt=i % 2 == 1)
        spacer(doc, 80)

    # Stats
    if assessment.get("stats"):
        st  = assessment["stats"]
        COL = [2400, 2400]
        tbl = doc.add_table(rows=0, cols=2)
        _set_table_width(tbl, sum(COL))
        _set_column_widths(tbl, COL)
        for i, (label, val) in enumerate([
            ("Total Incidents",    str(st.get("total", 0))),
            ("CCE %",             f"{st.get('cce','0.0')}%"),
            ("CSE %",             f"{st.get('cse','0.0')}%"),
            ("Under Investigation",str(st.get("ui", 0))),
        ]):
            r = tbl.add_row()
            d_cell(r, label, COL[0], alt=i % 2 == 1)
            d_cell(r, val,   COL[1], bold=True, alt=i % 2 == 1)
        spacer(doc, 80)

    # General assessment
    if assessment.get("general_assessment"):
        p = doc.add_paragraph()
        _para_spacing(p, before=160, after=80)
        run = p.add_run("General Assessment")
        _run_font(run, 22, bold=True, color_hex=DARK)
        body_text(doc, assessment["general_assessment"])
        spacer(doc, 80)

    # Per-crime sections
    for crime in assessment.get("per_crime") or []:
        # Crime heading with left border
        p = doc.add_paragraph()
        _para_spacing(p, before=240, after=80)
        run = p.add_run(crime.get("crime_type") or "")
        _run_font(run, 24, bold=True, color_hex=DARK)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"),   "single")
        left.set(qn("w:sz"),    "16")
        left.set(qn("w:space"), "4")
        left.set(qn("w:color"), DARK)
        pBdr.append(left)
        pPr.append(pBdr)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "120")
        pPr.append(ind)

        # Croston forecast
        croston_entry = None
        if analysis_data and analysis_data.get("croston", {}).get("per_crime"):
            croston_entry = next(
                (c for c in analysis_data["croston"]["per_crime"]
                 if c.get("crime") == crime.get("crime_type")),
                None,
            )
        if croston_entry:
            trend = croston_entry.get("trend")
            trend_label = (
                "↑ Increasing"       if trend == "increasing"        else
                "↓ Decreasing"       if trend == "decreasing"        else
                "Insufficient Data"  if trend == "insufficient_data" else
                "→ Stable"
            )
            pred = croston_entry.get("predicted_next_week")
            conf = croston_entry.get("confidence", 0)
            forecast_text = (
                f"{pred} incidents next week ({conf}% confidence)"
                if pred is not None
                else "Insufficient data for forecast"
            )
            COL = [2400, 8066]
            tbl = doc.add_table(rows=0, cols=2)
            _set_table_width(tbl, sum(COL))
            _set_column_widths(tbl, COL)
            r1 = tbl.add_row()
            d_cell(r1, "Trend",    COL[0], alt=False)
            d_cell(r1, trend_label,COL[1], bold=True, alt=False)
            r2 = tbl.add_row()
            d_cell(r2, "Forecast",    COL[0], alt=True)
            d_cell(r2, forecast_text, COL[1], alt=True)
            spacer(doc, 80)

        # QUAD sections
        quad_sections = [
            ("Crime Assessment",         crime.get("general_assessment")),
            ("Operations",               crime.get("operations")),
            ("Intelligence",             crime.get("intelligence")),
            ("Investigations",           crime.get("investigations")),
            ("Police Community Relations",crime.get("police_community_relations")),
        ]
        for label, value in quad_sections:
            if not value:
                continue
            p = doc.add_paragraph()
            _para_spacing(p, before=120, after=40)
            run = p.add_run(label)
            _run_font(run, 20, bold=True, color_hex="374151")

            if label == "Operations":
                for line in (l for l in value.split("\n") if l.strip()):
                    lp = doc.add_paragraph()
                    _para_spacing(lp, after=40)
                    lr = lp.add_run(line.replace("**", ""))
                    _run_font(lr, 18)
            else:
                vp = doc.add_paragraph()
                _para_spacing(vp, after=80)
                vr = vp.add_run(value)
                _run_font(vr, 18)

        spacer(doc, 80)


# ── Document builder ──────────────────────────────────────────────────────────
def build_export_doc(
    summary:       list,
    by_day:        list,
    place:         list,
    barangay:      list,
    modus:         list,
    complete_data: list,
    meta:          dict,
    images:        dict,
    assessment,
    analysis_data,
) -> bytes:
    totals = {
        "total":   sum(r["total"]              for r in summary),
        "cleared": sum(r["cleared"]            for r in summary),
        "solved":  sum(r["solved"]             for r in summary),
        "ui":      sum(r["underInvestigation"] for r in summary),
    }
    cce = pct(totals["cleared"] + totals["solved"], totals["total"])
    cse = pct(totals["solved"], totals["total"])

    now_str = (datetime.now(timezone.utc) + timedelta(hours=8)).strftime(
        "%B %d, %Y %I:%M %p"
    )

    doc = Document()

    # ── Page setup (A4 portrait) ───────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Twips(11906)
    section.page_height = Twips(16838)
    section.left_margin = section.right_margin = Twips(720)
    section.top_margin  = section.bottom_margin = Twips(720)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)

    # ── Cover page ─────────────────────────────────────────────────────────────
    for text, size, color, spacing_after in [
        ("CRIME DASHBOARD REPORT",  40, DARK,    60),
        ("Index Crime Statistics",  26, "6B7280", 60),
        (f"Reporting Period: {fmt_date_iso(meta.get('dateFrom')) or 'All dates'} — "
         f"{fmt_date_iso(meta.get('dateTo')) or 'All dates'}", 22, "000000", 40),
        (f"Crime Types: {', '.join(meta.get('crimeTypes') or []) or 'All Index Crimes'}",
         22, "000000", 40),
        (f"Barangays: {', '.join(meta.get('barangays') or []) or 'All Barangays'}",
         22, "000000", 40),
        (f"Generated: {now_str}", 20, "9CA3AF", 200),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(p, after=spacing_after)
        run = p.add_run(text)
        _run_font(run, size, bold=(size == 40), color_hex=color)

    # Horizontal rule
    hrule = doc.add_paragraph()
    _para_spacing(hrule, after=200)
    _para_border_bottom(hrule, color=DARK, size=8)

    # 1. Summary Statistics
    section_heading(doc, "1. Summary Statistics")
    build_summary_stat_table(doc, totals, cce, cse)
    spacer(doc)

    # 2. Index Crime Summary Table
    section_heading(doc, "2. Index Crime Summary Table")
    body_text(doc,
        "CCE (Crime Clearance Efficiency) = (Cleared + Solved) / Total  ·  "
        "CSE (Crime Solution Efficiency) = Solved / Total"
    )
    spacer(doc, 60)
    build_index_crime_table(doc, summary)
    spacer(doc)

    # 3. Crime Index Trends
    section_heading(doc, "3. Crime Index Trends")
    image_block(doc, images.get("trends"))
    spacer(doc, 60)

    # 4. Crime Clock
    section_heading(doc, "4. Crime Clock — Hourly Distribution")
    image_block(doc, images.get("clock"))
    spacer(doc, 60)

    # 5. Crime by Day of Week
    section_heading(doc, "5. Crime by Day of Week")
    build_by_day_table(doc, by_day)
    spacer(doc)

    # 6. Modus Operandi
    section_heading(doc, "6. Modus Operandi")
    build_modus_table(doc, modus)
    spacer(doc)

    # 7. Place of Commission
    section_heading(doc, "7. Place of Commission")
    build_place_table(doc, place)
    spacer(doc)

    # 8. Barangay Incidents
    section_heading(doc, "8. Barangay Incidents")
    build_barangay_table(doc, barangay)
    spacer(doc)

    # 9. Complete Data
    section_heading(doc, "9. Complete Data")
    body_text(doc, f"{len(complete_data)} total records")
    spacer(doc, 60)
    if complete_data:
        build_complete_data_table(doc, complete_data)
        spacer(doc)
    else:
        body_text(doc, "No records found for the selected filters.")

    # 10. AI Crime Assessment
    build_assessment_section(doc, assessment, analysis_data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── LibreOffice conversion ────────────────────────────────────────────────────
def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Writes docx_bytes to a temp file, calls LibreOffice headless to convert
    it to PDF, reads and returns the PDF bytes.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "report.docx")
        pdf_path  = os.path.join(tmpdir, "report.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, docx_path,
            ],
            capture_output=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr.decode()}"
            )

        with open(pdf_path, "rb") as f:
            return f.read()


# ── FastAPI handler ───────────────────────────────────────────────────────────
async def export_dashboard(req: Request):
    try:
        body = await req.json()

        summary       = body.get("summary",      [])
        by_day        = body.get("byDay",         [])
        place         = body.get("place",         [])
        barangay      = body.get("barangay",      [])
        modus         = body.get("modus",         [])
        complete_data = body.get("completeData",  [])
        meta          = body.get("meta",          {})
        images        = body.get("images",        {})
        assessment    = body.get("assessment",    None)
        analysis_data = body.get("analysisData",  None)

        user = getattr(req.state, "user", {})

        # Step 1: Build .docx
        docx_bytes = build_export_doc(
            summary, by_day, place, barangay, modus,
            complete_data, meta, images, assessment, analysis_data,
        )

        # Step 2: Convert to PDF via LibreOffice
        pdf_bytes = docx_to_pdf(docx_bytes)

        date_from = meta.get("dateFrom")
        date_to   = meta.get("dateTo")
        date_str  = (
            f"{date_from}_to_{date_to}"
            if date_from and date_to
            else datetime.now(timezone.utc).strftime("%Y-%m-%d")
        )

        await log_audit(
            user_id     = user.get("user_id"),
            username    = user.get("username"),
            event_name  = "Dashboard Export",
            description = f"Crime dashboard report exported ({date_str})",
            action      = "EXPORT",
            status      = "success",
            source      = "Web Portal",
            ip_address  = get_client_ip(req),
        )

        return Response(
            content     = pdf_bytes,
            media_type  = "application/pdf",
            headers     = {
                "Content-Disposition":
                    f'attachment; filename="crime_dashboard_{date_str}.pdf"'
            },
        )

    except Exception as e:
        print(f"export_dashboard error: {e}")
        user = getattr(req.state, "user", {})
        await log_audit(
            user_id     = user.get("user_id"),
            username    = user.get("username"),
            event_name  = "Dashboard Export Failed",
            description = str(e),
            action      = "EXPORT",
            status      = "failed",
            source      = "Web Portal",
            ip_address  = get_client_ip(req),
        )
        return JSONResponse(
            status_code = 500,
            content     = {"success": False, "message": str(e)},
        )